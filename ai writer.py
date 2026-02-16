import sys
import requests
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QLabel, QPushButton, QTextEdit, 
                             QComboBox, QMessageBox, QFrame, QSlider, QFileDialog, 
                             QToolBar, QStatusBar, QSplitter, QSizePolicy, QSpinBox)
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QSize
from PyQt5.QtGui import QFont, QTextCursor

# Try to import docx for .doc support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- Configuration ---
OLLAMA_URL = "http://localhost:11434"
DEFAULT_TOKEN_LIMIT = 140
DEFAULT_TEMPERATURE = 0.7

# --- Stylesheets ---
STYLES = {
    "light": """
        QMainWindow { background-color: #f5f5f7; }
        QWidget { font-family: 'Segoe UI', 'Roboto', 'Helvetica', sans-serif; color: #333; }
        
        QToolBar { 
            background-color: #ffffff; 
            border-bottom: 1px solid #e0e0e0; 
            padding: 8px; 
            spacing: 10px;
        }
        QToolBar QPushButton {
            background-color: transparent;
            border: 1px solid #ddd;
            border-radius: 6px;
            padding: 8px 15px;
            font-weight: 500;
            color: #333;
        }
        QToolBar QPushButton:hover { background-color: #e5e5ea; }
        QToolBar QPushButton:disabled { color: #ccc; border-color: #eee; }
        
        QToolBar QPushButton#generate-btn {
            background-color: #007aff;
            color: white;
            border: none;
            font-weight: 600;
        }
        QToolBar QPushButton#generate-btn:hover { background-color: #0056b3; }
        
        QTextEdit#editor {
            background-color: #ffffff;
            border: none;
            border-radius: 0px;
            padding: 40px;
            font-size: 16px;
            line-height: 1.6;
            font-family: 'Georgia', 'Times New Roman', serif;
        }
        QTextEdit#editor:focus { outline: none; }
        
        QFrame#sidebar {
            background-color: #ffffff;
            border-left: 1px solid #e0e0e0;
            padding: 20px;
        }
        
        QComboBox, QSpinBox {
            background-color: #f9f9f9;
            border: 1px solid #ddd;
            border-radius: 6px;
            padding: 8px;
            font-size: 13px;
        }
        QComboBox:focus, QSpinBox:focus { border: 1px solid #007aff; }
        
        QSlider::groove:horizontal { border: 1px solid #ddd; height: 6px; background: #e0e0e0; border-radius: 3px; }
        QSlider::handle:horizontal { background: #007aff; border: 1px solid #0056b3; width: 16px; margin: -5px 0; border-radius: 8px; }
        QSlider::handle:horizontal:hover { background: #0056b3; }
        
        QLabel#sidebar-title { font-size: 14px; font-weight: 600; color: #555; margin: 10px 0; }
        
        QStatusBar { background-color: #ffffff; border-top: 1px solid #e0e0e0; color: #888; }
        
        QPushButton#save-btn {
            background-color: #34c759;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 8px 15px;
            font-weight: 500;
            margin: 5px 0;
        }
        QPushButton#save-btn:hover { background-color: #28a745; }
        QPushButton#save-btn:disabled { background-color: #ccc; }
    """,
    "dark": """
        QMainWindow { background-color: #1e1e1e; }
        QWidget { font-family: 'Segoe UI', 'Roboto', 'Helvetica', sans-serif; color: #e0e0e0; }
        
        QToolBar { 
            background-color: #2c2c2c; 
            border-bottom: 1px solid #3a3a3a; 
            padding: 8px; 
            spacing: 10px;
        }
        QToolBar QPushButton {
            background-color: transparent;
            border: 1px solid #444;
            border-radius: 6px;
            padding: 8px 15px;
            font-weight: 500;
            color: #e0e0e0;
        }
        QToolBar QPushButton:hover { background-color: #3a3a3a; }
        QToolBar QPushButton:disabled { color: #555; border-color: #333; }
        
        QToolBar QPushButton#generate-btn {
            background-color: #0a84ff;
            color: white;
            border: none;
            font-weight: 600;
        }
        QToolBar QPushButton#generate-btn:hover { background-color: #409cff; }
        
        QTextEdit#editor {
            background-color: #252526;
            border: none;
            border-radius: 0px;
            padding: 40px;
            font-size: 16px;
            line-height: 1.6;
            font-family: 'Georgia', 'Times New Roman', serif;
            color: #d4d4d4;
        }
        QTextEdit#editor:focus { outline: none; }
        
        QFrame#sidebar {
            background-color: #2c2c2c;
            border-left: 1px solid #3a3a3a;
            padding: 20px;
        }
        
        QComboBox, QSpinBox {
            background-color: #3a3a3a;
            border: 1px solid #444;
            border-radius: 6px;
            padding: 8px;
            font-size: 13px;
            color: #fff;
        }
        QComboBox:focus, QSpinBox:focus { border: 1px solid #0a84ff; }
        
        QSlider::groove:horizontal { border: 1px solid #444; height: 6px; background: #3a3a3a; border-radius: 3px; }
        QSlider::handle:horizontal { background: #0a84ff; border: 1px solid #0056b3; width: 16px; margin: -5px 0; border-radius: 8px; }
        QSlider::handle:horizontal:hover { background: #409cff; }
        
        QLabel#sidebar-title { font-size: 14px; font-weight: 600; color: #aaa; margin: 10px 0; }
        
        QStatusBar { background-color: #2c2c2c; border-top: 1px solid #3a3a3a; color: #888; }
        
        QPushButton#save-btn {
            background-color: #34c759;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 8px 15px;
            font-weight: 500;
            margin: 5px 0;
        }
        QPushButton#save-btn:hover { background-color: #28a745; }
        QPushButton#save-btn:disabled { background-color: #555; }
    """
}

class OllamaWorker(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    models_loaded = pyqtSignal(list)

    def __init__(self, endpoint, model=None, prompt=None, temperature=0.7, token_limit=140):
        super().__init__()
        self.endpoint = endpoint
        self.model = model
        self.prompt = prompt
        self.temperature = temperature
        self.token_limit = token_limit

    def run(self):
        try:
            if self.endpoint == "scan":
                response = requests.get(f"{OLLAMA_URL}/api/tags", timeout=5)
                if response.status_code == 200:
                    data = response.json()
                    models = [m['name'] for m in data.get('models', [])]
                    self.models_loaded.emit(models)
                else:
                    self.error.emit(f"API Error: {response.status_code}")
            
            elif self.endpoint == "generate":
                system_instruction = "Continue writing from where the text ends naturally. Do not repeat what was already written. Do not add explanations, comments, or meta-text. Just continue the story or sentence seamlessly."
                full_prompt = f"{system_instruction}\n\nText to complete:\n{self.prompt}"
                
                payload = {
                    "model": self.model,
                    "prompt": full_prompt,
                    "stream": False,
                    "options": {
                        "num_predict": self.token_limit,
                        "temperature": self.temperature
                    }
                }
                response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=120)
                if response.status_code == 200:
                    data = response.json()
                    completion = data.get('response', '')
                    cleaned = self.clean_completion(self.prompt, completion)
                    self.finished.emit(cleaned)
                else:
                    self.error.emit(f"Generation Error: {response.status_code}")
        except requests.exceptions.ConnectionError:
            self.error.emit("Cannot connect to Ollama. Is it running?")
        except Exception as e:
            self.error.emit(str(e))
    
    def clean_completion(self, original, completion):
        original_stripped = original.strip()
        completion_stripped = completion.strip()
        
        if completion_stripped.startswith(original_stripped):
            completion_stripped = completion_stripped[len(original_stripped):].strip()
        
        prefixes = [
            "here's the continuation:", "continuation:", "continued:", 
            "here is the completion:", "here's the completion:",
            "the continuation is:", "completion:"
        ]
        for prefix in prefixes:
            if completion_stripped.lower().startswith(prefix):
                completion_stripped = completion_stripped[len(prefix):].strip()
        
        if completion_stripped.startswith('"') and not original_stripped.endswith('"'):
            completion_stripped = completion_stripped[1:].strip()
        
        return completion_stripped

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Writer")
        self.setMinimumSize(1000, 700)
        self.current_theme = "light"
        self.temperature = DEFAULT_TEMPERATURE
        self.token_limit = DEFAULT_TOKEN_LIMIT
        self.current_file = None
        self.setStyleSheet(STYLES[self.current_theme])

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        splitter = QSplitter(Qt.Horizontal)
        
        # --- Main Editor Area ---
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        editor_layout.setSpacing(0)
        editor_layout.setContentsMargins(0, 0, 0, 0)
        
        # Toolbar
        self.toolbar = QToolBar()
        self.toolbar.setMovable(False)
        self.toolbar.setIconSize(QSize(20, 20))
        
        self.theme_btn = QPushButton("🌙")
        self.theme_btn.setFixedWidth(40)
        self.theme_btn.clicked.connect(self.toggle_theme)
        self.toolbar.addWidget(self.theme_btn)
        
        self.toolbar.addSeparator()
        
        self.model_combo = QComboBox()
        self.model_combo.setMinimumWidth(180)
        self.model_combo.addItem("Select model...")
        self.toolbar.addWidget(QLabel("  Model: "))
        self.toolbar.addWidget(self.model_combo)
        
        self.scan_btn = QPushButton("🔄")
        self.scan_btn.setToolTip("Refresh Models")
        self.scan_btn.clicked.connect(self.scan_models)
        self.toolbar.addWidget(self.scan_btn)
        
        self.toolbar.addSeparator()
        
        self.generate_btn = QPushButton("✨ Generate")
        self.generate_btn.setObjectName("generate-btn")
        self.generate_btn.clicked.connect(self.start_generation)
        self.generate_btn.setEnabled(False)
        self.toolbar.addWidget(self.generate_btn)
        
        self.toolbar.addSeparator()
        
        self.save_txt_btn = QPushButton("📄 Save .txt")
        self.save_txt_btn.clicked.connect(self.save_txt)
        self.toolbar.addWidget(self.save_txt_btn)
        
        if DOCX_AVAILABLE:
            self.save_doc_btn = QPushButton("📕 Save .docx")
            self.save_doc_btn.clicked.connect(self.save_docx)
            self.toolbar.addWidget(self.save_doc_btn)
        else:
            self.toolbar.addWidget(QLabel("  (install python-docx for .docx)"))
        
        stretch_widget = QWidget()
        stretch_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.toolbar.addWidget(stretch_widget)
        
        self.char_label = QLabel("0 chars")
        self.toolbar.addWidget(self.char_label)
        
        editor_layout.addWidget(self.toolbar)
        
        # Main Editor
        self.editor = QTextEdit()
        self.editor.setObjectName("editor")
        self.editor.setPlaceholderText("Start writing your story here... Press 'Generate' to let AI continue from your cursor position.")
        self.editor.textChanged.connect(self.on_text_changed)
        editor_layout.addWidget(self.editor)
        
        splitter.addWidget(editor_widget)
        
        # --- Sidebar ---
        sidebar = QFrame()
        sidebar.setObjectName("sidebar")
        sidebar.setFixedWidth(280)
        sidebar_layout = QVBoxLayout(sidebar)
        
        # Temperature Control
        temp_title = QLabel("🌡️ Temperature")
        temp_title.setObjectName("sidebar-title")
        sidebar_layout.addWidget(temp_title)
        
        self.temp_value_label = QLabel(f"{self.temperature:.2f}")
        self.temp_value_label.setAlignment(Qt.AlignRight)
        self.temp_value_label.setStyleSheet("font-weight: bold; font-size: 16px; color: #007aff;")
        sidebar_layout.addWidget(self.temp_value_label)
        
        self.temp_slider = QSlider(Qt.Horizontal)
        self.temp_slider.setMinimum(0)
        self.temp_slider.setMaximum(200)
        self.temp_slider.setValue(int(DEFAULT_TEMPERATURE * 100))
        self.temp_slider.setTickPosition(QSlider.TicksBelow)
        self.temp_slider.setTickInterval(25)
        self.temp_slider.valueChanged.connect(self.on_temperature_changed)
        sidebar_layout.addWidget(self.temp_slider)
        
        temp_hints = QHBoxLayout()
        temp_hints.addWidget(QLabel("🎯"))
        temp_hints.addStretch()
        temp_hints.addWidget(QLabel("⚖️"))
        temp_hints.addStretch()
        temp_hints.addWidget(QLabel("🎨"))
        sidebar_layout.addLayout(temp_hints)
        
        sidebar_layout.addSpacing(20)
        
        # Token Limit Control (NEW)
        token_title = QLabel("📊 Token Limit")
        token_title.setObjectName("sidebar-title")
        sidebar_layout.addWidget(token_title)
        
        self.token_spinbox = QSpinBox()
        self.token_spinbox.setMinimum(10)
        self.token_spinbox.setMaximum(2000)
        self.token_spinbox.setValue(DEFAULT_TOKEN_LIMIT)
        self.token_spinbox.setSuffix(" tokens")
        self.token_spinbox.valueChanged.connect(self.on_token_limit_changed)
        sidebar_layout.addWidget(self.token_spinbox)
        
        token_hints = QLabel("Higher = longer response\nLower = shorter response")
        token_hints.setStyleSheet("font-size: 11px; color: #888;")
        sidebar_layout.addWidget(token_hints)
        
        sidebar_layout.addSpacing(20)
        
        # Info
        info_title = QLabel("ℹ️ Tips")
        info_title.setObjectName("sidebar-title")
        sidebar_layout.addWidget(info_title)
        
        tips = [
            "• Place cursor where you want AI to continue",
            "• Lower temp = more focused",
            "• Higher temp = more creative",
            "• Save your work frequently"
        ]
        for tip in tips:
            sidebar_layout.addWidget(QLabel(tip))
        
        sidebar_layout.addStretch()
        
        splitter.addWidget(sidebar)
        main_layout.addWidget(splitter)
        
        # Status Bar
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("Ready")
        
        self.scan_models()

    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.theme_btn.setText("☀️")
        else:
            self.current_theme = "light"
            self.theme_btn.setText("🌙")
        self.setStyleSheet(STYLES[self.current_theme])

    def on_temperature_changed(self, value):
        self.temperature = value / 100.0
        self.temp_value_label.setText(f"{self.temperature:.2f}")

    def on_token_limit_changed(self, value):
        self.token_limit = value
        self.statusBar.showMessage(f"Token limit set to {self.token_limit} tokens")

    def on_text_changed(self):
        text = self.editor.toPlainText()
        self.char_label.setText(f"{len(text)} chars")
        
        model = self.model_combo.currentText()
        has_text = len(text.strip()) > 0
        has_model = model not in ["Select model...", "No models found"]
        self.generate_btn.setEnabled(has_text and has_model)

    def scan_models(self):
        self.statusBar.showMessage("Scanning for models...")
        self.model_combo.clear()
        self.worker = OllamaWorker(endpoint="scan")
        self.worker.models_loaded.connect(self.on_models_loaded)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_models_loaded(self, models):
        self.model_combo.clear()
        if not models:
            self.model_combo.addItem("No models found")
            self.statusBar.showMessage("❌ No models available")
            self.generate_btn.setEnabled(False)
        else:
            self.model_combo.addItems(models)
            self.statusBar.showMessage(f"✓ {len(models)} models available")
            self.on_text_changed()

    def start_generation(self):
        text = self.editor.toPlainText()
        model = self.model_combo.currentText()
        
        if not text.strip() or model in ["Select model...", "No models found"]:
            QMessageBox.warning(self, "Warning", "Please enter text and select a model.")
            return

        self.generation_cursor_pos = len(text)
        
        self.generate_btn.setEnabled(False)
        self.generate_btn.setText("⏳ Generating...")
        self.statusBar.showMessage(f"AI is writing (Temp: {self.temperature:.2f}, Tokens: {self.token_limit})...")
        
        self.worker = OllamaWorker(endpoint="generate", model=model, prompt=text, 
                                   temperature=self.temperature, token_limit=self.token_limit)
        self.worker.finished.connect(self.on_generation_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_generation_finished(self, completion):
        if not completion.strip():
            self.statusBar.showMessage("⚠️ No completion generated")
            self.generate_btn.setEnabled(True)
            self.generate_btn.setText("✨ Generate")
            return
        
        cursor = self.editor.textCursor()
        cursor.setPosition(self.generation_cursor_pos)
        
        text = self.editor.toPlainText()
        if text and not text[-1].isspace() and completion and not completion[0].isspace():
            cursor.insertText(" ")
        
        cursor.insertText(completion)
        self.editor.setTextCursor(cursor)
        
        self.statusBar.showMessage(f"✓ Completion added (Temp: {self.temperature:.2f}, Tokens: {self.token_limit})")
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("✨ Generate")
        self.on_text_changed()

    def on_error(self, error_msg):
        self.statusBar.showMessage("❌ Error")
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("✨ Generate")
        QMessageBox.critical(self, "Error", error_msg)

    def save_txt(self):
        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Warning", "No text to save!")
            return
        
        default_name = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as TXT", default_name, "Text Files (*.txt)")
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.statusBar.showMessage(f"✓ Saved to {file_path}")
                self.current_file = file_path
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save: {str(e)}")

    def save_docx(self):
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx not installed. Run: pip install python-docx")
            return
        
        text = self.editor.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "Warning", "No text to save!")
            return
        
        default_name = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save as DOCX", default_name, "Word Documents (*.docx)")
        
        if file_path:
            try:
                doc = Document()
                doc.add_heading('AI Writer Document', 0)
                doc.add_paragraph(f'Created: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph()
                
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                
                doc.save(file_path)
                self.statusBar.showMessage(f"✓ Saved to {file_path}")
                self.current_file = file_path
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save: {str(e)}")

if __name__ == "__main__":
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())