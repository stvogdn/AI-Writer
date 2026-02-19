"""File management for AI Writer.

This module handles file operations including saving documents in various
formats and managing file paths.
"""

from datetime import datetime
from pathlib import Path

from PyQt5.QtWidgets import QFileDialog, QMessageBox, QWidget

# Try to import docx support
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ai_writer.config import get_settings


class FileManager:
    """Manages file operations for the AI Writer application."""

    def __init__(self, parent: QWidget | None = None):
        """Initialize the file manager.

        Args:
            parent: Parent widget for dialogs
        """
        self.parent = parent
        self.current_file: str | None = None

    def save_as_txt(self, text: str) -> bool:
        """Save text as a .txt file.

        Args:
            text: Text content to save

        Returns:
            True if saved successfully, False otherwise
        """
        if not text.strip():
            self._show_warning("No text to save!")
            return False

        file_path = self._get_save_path("txt", "Text Files (*.txt)")
        if not file_path:
            return False

        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(text)
            self.current_file = file_path
            self._show_success(f"Saved to {file_path}")
            return True
        except Exception as e:
            self._show_error(f"Failed to save: {str(e)}")
            return False

    def save_as_docx(self, text: str) -> bool:
        """Save text as a .docx file.

        Args:
            text: Text content to save

        Returns:
            True if saved successfully, False otherwise
        """
        if not DOCX_AVAILABLE:
            self._show_warning(
                "python-docx not installed. Run: pip install python-docx"
            )
            return False

        if not text.strip():
            self._show_warning("No text to save!")
            return False

        file_path = self._get_save_path("docx", "Word Documents (*.docx)")
        if not file_path:
            return False

        try:
            doc = Document()
            doc.add_heading("AI Writer Document", 0)
            doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph()

            # Split into paragraphs and add to document
            paragraphs = text.split("\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)

            doc.save(file_path)
            self.current_file = file_path
            self._show_success(f"Saved to {file_path}")
            return True
        except Exception as e:
            self._show_error(f"Failed to save: {str(e)}")
            return False

    def save_with_default_format(self, text: str) -> bool:
        """Save using the default format from settings.

        Args:
            text: Text content to save

        Returns:
            True if saved successfully, False otherwise
        """
        settings = get_settings()
        default_format = settings.file.default_save_format.lower()

        if default_format == "docx":
            return self.save_as_docx(text)
        else:
            return self.save_as_txt(text)

    def load_file(self, file_path: str) -> str | None:
        """Load text from a file.

        Args:
            file_path: Path to the file to load

        Returns:
            Text content if loaded successfully, None otherwise
        """
        try:
            path = Path(file_path)

            if path.suffix.lower() in [".txt", ".md"]:
                with open(file_path, encoding="utf-8") as f:
                    content = f.read()
            elif path.suffix.lower() == ".docx" and DOCX_AVAILABLE:
                doc = Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                self._show_error(f"Unsupported file format: {path.suffix}")
                return None

            self.current_file = file_path
            return content
        except Exception as e:
            self._show_error(f"Failed to load file: {str(e)}")
            return None

    def _get_save_path(self, extension: str, file_filter: str) -> str | None:
        """Get save path from user dialog.

        Args:
            extension: File extension (without dot)
            file_filter: Qt file filter string

        Returns:
            Selected file path or None if cancelled
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"document_{timestamp}.{extension}"

        if self.parent:
            file_path, _ = QFileDialog.getSaveFileName(
                self.parent, f"Save as {extension.upper()}", default_name, file_filter
            )
        else:
            file_path = ""  # Fallback for testing

        return file_path if file_path else None

    def _show_warning(self, message: str) -> None:
        """Show a warning message."""
        if self.parent:
            QMessageBox.warning(self.parent, "Warning", message)
        else:
            print(f"Warning: {message}")

    def _show_error(self, message: str) -> None:
        """Show an error message."""
        if self.parent:
            QMessageBox.critical(self.parent, "Error", message)
        else:
            print(f"Error: {message}")

    def _show_success(self, message: str) -> None:
        """Show a success message."""
        # This would typically update a status bar
        # For now, just print to console in non-GUI contexts
        if not self.parent:
            print(f"Success: {message}")

    @property
    def has_docx_support(self) -> bool:
        """Check if DOCX support is available."""
        return DOCX_AVAILABLE

    @property
    def current_file_path(self) -> str | None:
        """Get the current file path."""
        return self.current_file
